"""
handler.py

Core file‐watcher handlers for:
  - Indexing new files into the DB
  - Generating thumbnails for PDFs, images, TXT, and DOCX
  - Cleaning up on file deletion
  - Housekeeping (orphan cleanup)

Functions:
  - process_new_file(filepath)
  - process_deleted_file(filepath)
  - cleanup_orphans(watch_folder)
  - create_thumbnail(doc_id, filepath, filename) -> slug
"""

import os
import re
from datetime import datetime

import fitz  # PyMuPDF for PDF rendering
from PIL import Image, ImageDraw, ImageFont
from docx import Document as DocxDocument

from sqlalchemy.exc import SQLAlchemyError
from database import SessionLocal, Document, Thumbnail
from event_queue import event_queue

# Directory for storing thumbnails
THUMB_DIR = os.path.join("static", "thumbnails")
os.makedirs(THUMB_DIR, exist_ok=True)


def _slugify(name: str) -> str:
    """
    Convert an arbitrary filename base into a lowercase, underscore‐separated slug.

    Examples:
      "My File Name.docx" → "my_file_name"
      "Résumé 2025.pdf"   → "r_sum_2025"
    """
    return re.sub(r'[^0-9a-zA-Z]+', '_', name).strip('_').lower()


def process_new_file(filepath: str):
    """
    Handle a newly created or detected file.

    Steps:
      1) If not already in the DB, insert a new Document row.
      2) Generate a thumbnail (PNG) and store it in the thumbnails table.
      3) Emit a Server‐Sent Event ("created") with the thumbnail URL.

    Args:
      filepath: Full path to the file that was created.
    """
    session = SessionLocal()
    filename = os.path.basename(filepath)

    try:
        # 1) Skip if already indexed
        if session.query(Document).filter_by(filename=filename).first():
            print(f"⚠️  Already indexed: {filename}")
            return

        # 2) Insert the document record
        doc = Document(filename=filename, path=filepath, source="fallback")
        session.add(doc)
        session.commit()
        print(f"✅ Indexed document: {filename}")

        # 3) Create and record thumbnail; get the slug basename
        slug = create_thumbnail(doc.id, filepath, filename)

        # 4) Broadcast a 'created' SSE event
        event_queue.put({
            "action":    "created",
            "id":        doc.id,
            "filename":  filename,
            "thumbnail": f"/static/thumbnails/{slug}.png"
        })

    except SQLAlchemyError as e:
        session.rollback()
        print(f"❌ DB error indexing {filename}: {e}")
    except Exception as e:
        print(f"❌ Unexpected error indexing {filename}: {e}")
    finally:
        session.close()


def process_deleted_file(filepath: str):
    """
    Handle removal of a file from the watch folder.

    Steps:
      1) Delete any associated thumbnail files and DB rows.
      2) Delete the Document row.
      3) Emit a Server‐Sent Event ("deleted") with the filename.

    Args:
      filepath: Full path to the file that was deleted.
    """
    session = SessionLocal()
    filename = os.path.basename(filepath)
    try:
        doc = session.query(Document).filter_by(filename=filename).first()
        if not doc:
            print(f"⚠️  Not in DB: {filename}")
            return

        # 1) Delete all thumbnail files + DB rows
        for thumb in list(doc.thumbnails):
            try:
                os.remove(thumb.thumbnail_path)
                print(f"🗑️  Thumbnail removed: {thumb.thumbnail_path}")
            except OSError:
                pass
            session.delete(thumb)

        # 2) Delete document record
        session.delete(doc)
        session.commit()
        print(f"🗑️  Removed document & thumbnails: {filename}")

        # 3) Broadcast deletion event
        event_queue.put({"action": "deleted", "filename": filename})

    except Exception as e:
        session.rollback()
        print(f"❌ Error removing {filename}: {e}")
    finally:
        session.close()


def cleanup_orphans(watch_folder: str):
    """
    Remove any Document entries whose files no longer exist on disk.

    Args:
      watch_folder: Path to the folder being watched.
    """
    session = SessionLocal()
    try:
        for doc in session.query(Document).all():
            if not os.path.exists(doc.path):
                print(f"🧹 Orphan cleanup: {doc.filename}")
                process_deleted_file(doc.path)
    finally:
        session.close()


def create_thumbnail(doc_id: int, filepath: str, filename: str) -> str:
    """
    Generate a PNG thumbnail for `filepath`, save it under THUMB_DIR,
    insert a Thumbnail row, and return the slug basename.

    Supports:
      - PDFs  → first page at 300 DPI
      - Images (png/jpg/jpeg/gif) → resized to max 400×400
      - .txt  → first 20 lines rendered on an 800×600 canvas
      - .docx → first 5 paragraphs rendered on an 800×600 canvas
      - Others → generic ‘FILE’ placeholder

    Args:
      doc_id:   Primary‐key ID of the Document row.
      filepath: Full path to the source file.
      filename: The basename (with extension) of the file.

    Returns:
      slug: The URL‐safe slug used for naming `slug.png` under THUMB_DIR.
    """
    base, ext = os.path.splitext(filename)
    slug = _slugify(base)
    thumb_filename = f"{slug}.png"
    thumb_path     = os.path.join(THUMB_DIR, thumb_filename)

    try:
        ext_low = ext.lower()
        if ext_low == ".pdf":
            _generate_pdf_thumbnail(filepath, thumb_path)
        elif ext_low in (".png", ".jpg", ".jpeg", ".gif"):
            _generate_image_thumbnail(filepath, thumb_path)
        elif ext_low == ".txt":
            _generate_txt_thumbnail(filepath, thumb_path)
        elif ext_low == ".docx":
            _generate_docx_thumbnail(filepath, thumb_path)
        else:
            _copy_fallback_icon(thumb_path)

        # Record in thumbnails table
        session = SessionLocal()
        session.add(Thumbnail(document_id=doc_id,
                              thumbnail_path=thumb_path,
                              created_at=datetime.utcnow()))
        session.commit()
        session.close()

        print(f"🖼️  Created thumbnail for: {filename} → {thumb_filename}")
    except Exception as e:
        print(f"❌ Thumbnail error for {filename}: {e}")

    return slug


# ----- Renderer Helpers ----- #

def _generate_pdf_thumbnail(src: str, dst: str):
    """Render the first page of a PDF as a 300 DPI PNG."""
    doc  = fitz.open(src)
    page = doc.load_page(0)
    pix  = page.get_pixmap(dpi=300)
    pix.save(dst)
    doc.close()


def _generate_image_thumbnail(src: str, dst: str):
    """Resize an image to max 400×400 and save as PNG."""
    with Image.open(src) as img:
        img.thumbnail((400, 400))
        img.save(dst)


def _generate_txt_thumbnail(src: str, dst: str):
    """Render the first 20 lines of a .txt onto an 800×600 white canvas."""
    lines = []
    with open(src, 'r', encoding='utf-8', errors='ignore') as f:
        for _ in range(20):
            try:
                lines.append(next(f))
            except StopIteration:
                break
    text = ''.join(lines).strip() or 'Empty file'
    _render_text_to_image(text, dst)


def _generate_docx_thumbnail(src: str, dst: str):
    """Render the first 5 non-empty paragraphs of a .docx onto an 800×600 canvas."""
    doc   = DocxDocument(src)
    paras = [p.text for p in doc.paragraphs[:5] if p.text.strip()]
    text  = '\n'.join(paras) or 'No text'
    _render_text_to_image(text, dst)


def _render_text_to_image(text: str, dst: str, size=(800, 600)):
    """
    Helper to draw multiline text onto a white canvas.

    Args:
      text: The text to render (will be wrapped to 20 lines).
      dst:  Path to save the PNG.
      size:(width, height) of the canvas.
    """
    img  = Image.new("RGB", size, color="white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    bbox  = draw.textbbox((0, 0), "Hg", font=font)
    lh    = (bbox[3] - bbox[1]) + 4
    for i, line in enumerate(text.split("\n")[:20]):
        draw.text((10, 10 + i * lh), line, fill="black", font=font)
    img.save(dst)


def _copy_fallback_icon(dst: str):
    """Draw a simple ‘FILE’ placeholder on a gray square."""
    img  = Image.new("RGB", (200, 200), color=(240, 240, 240))
    draw = ImageDraw.Draw(img)
    draw.text((50, 90), "FILE", fill="black", font=ImageFont.load_default())
    img.save(dst)
